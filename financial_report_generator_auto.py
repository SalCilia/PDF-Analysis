#!/usr/bin/env python3
"""
Financial Data Extraction and Report Generation Script
With automatic Downloads folder detection for POW Funds Report
"""

import os
import sys
import time
import logging
import random
import glob
from typing import List, Tuple, Optional
from urllib.parse import quote

import pdfplumber
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('financial_report.log'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)


def find_pow_funds_pdf():
    """
    Automatically locate POW Funds Report PDF in Downloads folder.
    
    Returns:
        str: Path to the PDF file, or None if not found.
    """
    # Get Downloads folder path for current user
    downloads_path = os.path.join(os.path.expanduser("~"), "Downloads")
    
    logger.info(f"Searching for PDF files in: {downloads_path}")
    
    # Search for files containing "POW" in the name
    search_patterns = [
        os.path.join(downloads_path, "*POW*Funds*.pdf"),
        os.path.join(downloads_path, "*POW*funds*.pdf"),
        os.path.join(downloads_path, "*POW*.pdf"),
    ]
    
    for pattern in search_patterns:
        files = glob.glob(pattern, recursive=False)
        if files:
            pdf_file = files[0]
            logger.info(f"✓ Found POW Funds PDF: {pdf_file}")
            return pdf_file
    
    # If no pattern matched, list all PDFs in Downloads
    logger.warning("POW Funds PDF not found with specific patterns. Checking all PDFs...")
    all_pdfs = glob.glob(os.path.join(downloads_path, "*.pdf"))
    
    if all_pdfs:
        logger.info(f"Available PDF files in Downloads:")
        for idx, pdf in enumerate(all_pdfs, 1):
            logger.info(f"  {idx}. {os.path.basename(pdf)}")
    else:
        logger.error("❌ No PDF files found in Downloads folder")
        return None
    
    logger.error("❌ Could not locate POW Funds Report PDF")
    return None


class FinancialReportGenerator:
    """Handles extraction of holdings from PDF and generation of Word report."""

    CAUTIOUS_HEADERS = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
                      '(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept-Language': 'en-US,en;q=0.9',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,'
                  'image/webp,image/apng,*/*;q=0.8',
        'Accept-Encoding': 'gzip, deflate, br',
        'DNT': '1',
        'Connection': 'keep-alive',
        'Upgrade-Insecure-Requests': '1'
    }

    def __init__(self, pdf_file_path: str):
        self.pdf_file_path = pdf_file_path
        self.holdings = []
        self.search_results = {}

    def extract_top_holdings(self, num_holdings: int = 10) -> List[str]:
        """Extract top holdings from the PDF file."""
        try:
            if not os.path.exists(self.pdf_file_path):
                logger.error(f"PDF file not found: {self.pdf_file_path}")
                return []

            logger.info(f"Reading PDF file: {self.pdf_file_path}")

            with pdfplumber.open(self.pdf_file_path) as pdf:
                full_text = ""
                for page in pdf.pages:
                    full_text += page.extract_text() + "\n"

                holdings = self._parse_holdings_section(full_text, num_holdings)

                if holdings:
                    logger.info(f"✓ Extracted {len(holdings)} holdings from PDF")
                    self.holdings = holdings
                    return holdings
                else:
                    logger.warning("No holdings found in PDF")
                    return []

        except pdfplumber.PDFException as e:
            logger.error(f"Error reading PDF: {e}")
            return []
        except Exception as e:
            logger.error(f"Unexpected error: {e}")
            return []

    def _parse_holdings_section(self, text: str, num_holdings: int) -> List[str]:
        """Parse the holdings section from extracted PDF text."""
        holdings = []

        section_markers = [
            "top ten holdings",
            "top 10 holdings",
            "top holding",
            "holdings",
            "portfolio holding",
            "fund holding",
        ]

        text_lower = text.lower()
        start_idx = -1

        for marker in section_markers:
            idx = text_lower.find(marker)
            if idx != -1:
                start_idx = idx
                logger.info(f"Found holdings section: '{marker}'")
                break

        if start_idx == -1:
            logger.warning("Could not find holdings section")
            return []

        section_text = text[start_idx:]
        lines = section_text.split('\n')

        for line in lines[1:]:
            line = line.strip()

            if not line or len(line) < 2:
                continue

            if any(skip in line.lower() for skip in ['total', 'fund', 'report', 'date', 'sector', 'cash', 'other']):
                continue

            company_name = self._clean_company_name(line)

            if company_name and company_name not in holdings:
                holdings.append(company_name)
                logger.debug(f"Extracted: {company_name}")

            if len(holdings) >= num_holdings:
                break

        return holdings[:num_holdings]

    @staticmethod
    def _clean_company_name(line: str) -> str:
        """Clean and extract company name."""
        company_name = line.split('%')[0].strip()
        company_name = company_name.split('Inc')[0].strip()
        company_name = company_name.split('Ltd')[0].strip()
        company_name = company_name.split('Corp')[0].strip()
        company_name = company_name.split('Co.')[0].strip()

        while company_name and (company_name[0].isdigit() or company_name[0] in '.-'):
            company_name = company_name[1:].strip()

        return company_name if len(company_name) > 1 else ""

    def search_companies(self, delay_range: Tuple[int, int] = (2, 5)) -> dict:
        """Perform web searches for each extracted company."""
        if not self.holdings:
            logger.error("No holdings to search")
            return {}

        logger.info(f"Starting web search for {len(self.holdings)} companies")

        for idx, company in enumerate(self.holdings):
            logger.info(f"[{idx + 1}/{len(self.holdings)}] Searching: {company}")

            results = self._search_company(company)
            self.search_results[company] = results

            if idx < len(self.holdings) - 1:
                delay = random.uniform(delay_range[0], delay_range[1])
                logger.debug(f"Waiting {delay:.2f}s before next search...")
                time.sleep(delay)

        logger.info(f"✓ Web search completed")
        return self.search_results

    def _search_company(self, company_name: str, num_results: int = 3) -> List[dict]:
        """Search for a company on Google."""
        search_results = []

        try:
            search_query = f"{company_name} company"
            search_url = f"https://www.google.com/search?q={quote(search_query)}"

            response = requests.get(
                search_url,
                headers=self.CAUTIOUS_HEADERS,
                timeout=10,
                allow_redirects=True
            )

            if response.status_code != 200:
                logger.warning(f"Google returned status {response.status_code}")
                return search_results

            soup = BeautifulSoup(response.content, 'html.parser')
            search_result_divs = soup.find_all('div', {'data-sokoban-container': True})

            if not search_result_divs:
                search_result_divs = soup.find_all('div', {'class': 'g'})

            result_count = 0
            for result_div in search_result_divs:
                if result_count >= num_results:
                    break

                try:
                    title_elem = result_div.find('h3')
                    if not title_elem:
                        continue

                    title = title_elem.get_text(strip=True)

                    snippet_elem = result_div.find('span', {'class': 'aCOpf'}) or \
                                   result_div.find('span', {'class': 's'}) or \
                                   result_div.find('div', {'class': 'VwiC3b'})

                    snippet = ""
                    if snippet_elem:
                        snippet = snippet_elem.get_text(strip=True)

                    if title and snippet:
                        search_results.append({'title': title, 'snippet': snippet})
                        result_count += 1

                except Exception as e:
                    logger.debug(f"Error parsing result: {e}")
                    continue

        except requests.RequestException as e:
            logger.error(f"Network error: {e}")
        except Exception as e:
            logger.error(f"Search error: {e}")

        return search_results

    def generate_word_report(self, output_file: str = "Financial_Holdings_Report.docx") -> bool:
        """Generate a Word document report."""
        try:
            logger.info("Generating Word document...")

            doc = Document()

            title = doc.add_heading('Financial Holdings Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitle = doc.add_paragraph(f'Top {len(self.holdings)} Holdings Analysis')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_format = subtitle.runs[0]
            subtitle_format.font.size = Pt(12)
            subtitle_format.font.italic = True

            timestamp = doc.add_paragraph()
            timestamp.add_run(f'Generated: {time.strftime("%Y-%m-%d %H:%M:%S")}').font.size = Pt(10)

            source_info = doc.add_paragraph()
            source_info.add_run(f'Source: {os.path.basename(self.pdf_file_path)}').font.size = Pt(10)

            doc.add_paragraph()

            for company in self.holdings:
                heading = doc.add_heading(company, level=1)
                heading_format = heading.runs[0]
                heading_format.font.color.rgb = RGBColor(0, 51, 102)

                results = self.search_results.get(company, [])

                if results:
                    for idx, result in enumerate(results, 1):
                        para = doc.add_paragraph(style='List Bullet')

                        title_run = para.add_run(f"Result {idx}: {result['title']}\n")
                        title_run.bold = True
                        title_run.font.size = Pt(11)

                        snippet_run = para.add_run(result['snippet'])
                        snippet_run.font.size = Pt(10)
                        snippet_run.font.italic = True
                else:
                    para = doc.add_paragraph(style='List Bullet')
                    no_results_run = para.add_run("No search results available")
                    no_results_run.font.italic = True

                doc.add_paragraph()

            doc.save(output_file)
            logger.info(f"✓ Word document saved: {output_file}")
            return True

        except Exception as e:
            logger.error(f"Error generating report: {e}")
            return False


def main():
    """Main function."""
    
    logger.info("=" * 80)
    logger.info("Financial Holdings Report Generator")
    logger.info("=" * 80)

    logger.info("\n[Step 0] Locating POW Funds Report PDF...")
    PDF_FILE = find_pow_funds_pdf()

    if not PDF_FILE:
        logger.error("\n❌ Could not find POW Funds Report PDF in Downloads folder")
        logger.error("Please ensure the PDF is in your Downloads folder and try again.")
        return False

    OUTPUT_FILE = "Financial_Holdings_Report.docx"
    generator = FinancialReportGenerator(PDF_FILE)

    logger.info("\n[Step 1] Extracting holdings from PDF...")
    holdings = generator.extract_top_holdings()

    if not holdings:
        logger.error("❌ Failed to extract holdings from PDF")
        return False

    logger.info(f"Holdings: {', '.join(holdings)}\n")

    logger.info("[Step 2] Performing web searches...")
    generator.search_companies()

    logger.info("\n[Step 3] Generating Word document...")
    success = generator.generate_word_report(OUTPUT_FILE)

    if success:
        logger.info(f"\n✓ SUCCESS!")
        logger.info(f"Report saved as: {OUTPUT_FILE}")
        logger.info(f"PDF processed: {os.path.basename(PDF_FILE)}")
        logger.info("=" * 80 + "\n")
        return True
    else:
        logger.error("\n❌ Failed to generate report")
        return False


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)