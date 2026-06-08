#!/usr/bin/env python3
"""
Demo Mode - Simulates the application without needing a real PDF
This helps you understand how the application works
"""

import time
import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def demo_report_generation():
    """Generate a demo Word report with sample data."""
    
    logger.info("=" * 80)
    logger.info("DEMO MODE - Financial Holdings Report Generator")
    logger.info("=" * 80)
    
    # Demo data (simulating extracted holdings and search results)
    demo_holdings = {
        'Apple Inc.': [
            {
                'title': 'Apple - Wikipedia',
                'snippet': 'Apple Inc. is an American multinational technology company that designs and manufactures electronics and computer software.'
            },
            {
                'title': 'Apple Official Website',
                'snippet': 'Discover the latest Apple products and services. Shop iPhone, iPad, Apple Watch, Mac and more.'
            },
            {
                'title': 'Apple Inc. - Investor Relations',
                'snippet': 'Apple financial information and investor relations, including quarterly earnings and annual reports.'
            }
        ],
        'Microsoft Corporation': [
            {
                'title': 'Microsoft - Official Website',
                'snippet': 'Cloud-based computing services from Microsoft. Office 365, Azure, Windows and more.'
            },
            {
                'title': 'Microsoft - Wikipedia',
                'snippet': 'Microsoft is an American technology corporation that develops software, consumer electronics and personal computers.'
            },
            {
                'title': 'Microsoft Investor Relations',
                'snippet': 'Microsoft financial performance, earnings reports, and shareholder information.'
            }
        ],
        'Tesla Inc.': [
            {
                'title': 'Tesla - Official Website',
                'snippet': 'Tesla designs and manufactures electric vehicles, solar and energy storage systems.'
            },
            {
                'title': 'Tesla - Wikipedia',
                'snippet': 'Tesla, Inc. is an American electric vehicle and clean energy company founded by Elon Musk.'
            },
            {
                'title': 'Tesla Investor Relations',
                'snippet': 'Tesla Inc. financial data, stock performance, and corporate information for investors.'
            }
        ],
        'Amazon.com Inc.': [
            {
                'title': 'Amazon - Official Website',
                'snippet': 'Amazon offers fast, reliable shopping with Prime membership. Shop millions of items with free delivery.'
            },
            {
                'title': 'Amazon - Wikipedia',
                'snippet': 'Amazon is an American multinational technology company that focuses on e-commerce and cloud computing.'
            },
            {
                'title': 'Amazon Investor Relations',
                'snippet': 'Amazon.com Inc. investor relations, quarterly earnings, and shareholder information.'
            }
        ],
        'Alphabet Inc. (Google)': [
            {
                'title': 'Google - Official Website',
                'snippet': 'Google is a search engine and technology company offering search, Gmail, Drive, and more services.'
            },
            {
                'title': 'Alphabet Inc. - Wikipedia',
                'snippet': 'Alphabet Inc. is an American multinational technology conglomerate that owns Google and other subsidiaries.'
            },
            {
                'title': 'Alphabet Investor Relations',
                'snippet': 'Alphabet Inc. financial information, earnings reports, and investor resources.'
            }
        ]
    }
    
    logger.info("\n[Step 1] Extracting Holdings from PDF...")
    logger.info(f"Found {len(demo_holdings)} sample holdings")
    for company in demo_holdings.keys():
        logger.info(f"  ✓ {company}")
        time.sleep(0.3)
    
    logger.info("\n[Step 2] Web Search Simulation...")
    for idx, company in enumerate(demo_holdings.keys(), 1):
        logger.info(f"[{idx}/{len(demo_holdings)}] Searching: {company}")
        time.sleep(1)  # Simulate search delay
    
    logger.info("\n[Step 3] Generating Word Document...")
    
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Financial Holdings Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph(f'Top {len(demo_holdings)} Holdings Analysis (DEMO MODE)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(12)
    subtitle_format.font.italic = True
    
    # Add timestamp
    import time as time_module
    timestamp = doc.add_paragraph()
    timestamp.add_run(f'Generated: {time_module.strftime("%Y-%m-%d %H:%M:%S")}').font.size = Pt(10)
    
    # Add demo note
    demo_note = doc.add_paragraph()
    demo_run = demo_note.add_run('Note: This is a DEMO with sample data. Real data requires an actual PDF.')
    demo_run.font.italic = True
    demo_run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph()
    
    # Add company data
    for company, results in demo_holdings.items():
        # Add company heading
        heading = doc.add_heading(company, level=1)
        heading_format = heading.runs[0]
        heading_format.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Add results
        for idx, result in enumerate(results, 1):
            para = doc.add_paragraph(style='List Bullet')
            
            title_run = para.add_run(f"Result {idx}: {result['title']}\n")
            title_run.bold = True
            title_run.font.size = Pt(11)
            
            snippet_run = para.add_run(result['snippet'])
            snippet_run.font.size = Pt(10)
            snippet_run.font.italic = True
        
        doc.add_paragraph()
    
    # Save document
    output_file = "Demo_Holdings_Report.docx"
    doc.save(output_file)
    logger.info(f"✓ Word document saved: {output_file}")
    
    logger.info("\n" + "=" * 80)
    logger.info("✓ DEMO COMPLETE!")
    logger.info(f"Generated: {output_file}")
    logger.info("=" * 80)
    logger.info("\nNow run the real application with your PDF:")
    logger.info("  python financial_report_generator_auto.py\n")

if __name__ == "__main__":
    demo_report_generation()